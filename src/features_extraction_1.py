       # -*- coding: utf-8 -*-
"""
Created on Sun Jan 10 18:54:22 2021

@author: KHC
"""
# -*- coding: utf-8 -*-
"""
Created on Sun Jan  3 18:42:34 2021

@author: KHC
"""
# -*- coding: utf-8 -*-
"""
Created on Fri Jan  1 00:08:47 2021

@author: KHC
"""
# import fitz
from operator import itemgetter
import fr_core_news_lg
import os
import spacy
import textract
from os import listdir
from os.path import isfile, join
import docx
import re
import csv
import pandas as pd
import os
import nltk
import datetime
from dateparser.search import search_dates
from calendar import month_name,month_abbr
from datetime import date
# from sklearn.feature_extraction.text import CountVectorizer
from itertools import islice
import nltk
import pickle
from nltk.tokenize import word_tokenize
import numpy as np
from nltk.corpus import stopwords
from nltk import download
from nltk.stem import WordNetLemmatizer
from nltk.stem import PorterStemmer 
 
from natsort import natsorted
from nltk.stem import PorterStemmer 
from datasketch import MinHash, MinHashLSH
from nltk import ngrams
from nltk import flatten
import pickle

# nltk.download('punkt')
# download('stopwords')
lemmatizer = WordNetLemmatizer() 
import unidecode

from docx import Document
from datetime import datetime
from docx import Document

import glob
import win32com.client



import docx2txt

from utils_cvs import top_frequent

# this function will take a single CV and will return a list of sentences extracted
#from that CV. It will extract paragraphs as first step and then will convert 
# paragraphs into sentences and then will make a list of sentences. 

encoding = "utf-8"
def convert_to_paras(filename):

    try:
        doc=docx.Document(filename)
        all_paras = doc
    # #         print('len',len(doc.paragraphs))
    # #         print(all_paras)
        if len(doc.paragraphs)<=1:
            my_text = docx2txt.process(filename)
            d=Document()
            d.add_paragraph(my_text)
    
            all_paras = Document() 
    # #             print(all_paras)
            for s in d.paragraphs:
                str_list = s.text.split("\n")
                str_list[:] = [x.strip() for x in str_list if x]
                new_para = ""
                prev = str_list[0]
                for i in range(1, len(str_list)):
                    if new_para =="":
                        new_para = new_para + prev
                    else:
                        new_para = new_para + "\n" + prev
                    prev = str_list[i]
     #                         print('%%%%%%%%%%%%%','in')
                    all_paras.add_paragraph(prev)
     #         print('after',all_paras.paragraphs)
        return all_paras
    except:
         return "" 
        
       
def extract_qualifications(i,sentence,qualifications_list,cv_name):
#    lsm_q = spacy.load(r"C:\Projects\CVsRnkingAmir\code with dataset\code with dataset\model\Qualification_model_fre")
    doc_qual=lsm_q(sentence)
    for X in doc_qual.ents:
        if X.label_=="qualification":
            contains_digit = any(map(str.isdigit, X.text))
            if contains_digit or ':' in X.text:
#                print ('it contains numeric value')
                pass
            else:
            
                qualifications_list.append(X.text)
               

    qualifications_list=top_frequent(qualifications_list, k=15)
#    print ('length of qualification list is:',len(qualifications_list))
        
    return(qualifications_list)
            
def extract_skills(i,sentence,skills_list,cv_name):
#    lsm_s = spacy.load(r"C:\Projects\CVsRnkingAmir\code with dataset\code with dataset\model\skill_model_fre")
    doc_skills=lsm_s(sentence)
    for X in doc_skills.ents:
        if X.label_=="skills":
            contains_digit = any(map(str.isdigit, X.text))
            if contains_digit or ':' in X.text:
#                print ('it contains numeric value')
                pass
            else:
#                print ('it does not contain numeric value')
                skills_list.append(X.text)
    skills_list=top_frequent(skills_list, k=15)
#    print ('length of skills list is:',len(skills_list))
    return(skills_list)
        
            
def extract_locations(i,sentence,locations_list,cv_name):
    global nlp,l_list
    doc_all=nlp(sentence)
    for X in doc_all.ents:
        if (X.label_=='GPE'):
            locations_list.append(X.text)
            
    words_list=sentence.split(' ')

    for w in words_list:
        if w in l_list:
            # print ('found location')
            locations_list.append(w)
    loc_list=top_frequent(locations_list, k=3)
#    print ('length of locations list is:',len(loc_list))
    return(loc_list)
            
def extract_languages(i,sentence,languages_list,cv_name):
    global nlp,lang_list
    doc_all=nlp(sentence)
    for X in doc_all.ents:
        if (X.label_=='LANGUAGE'):
            languages_list.append(X.text)
            
            
    words_list=sentence.split(' ')

    for w in words_list:
        if w in lang_list:
            # print ('found language')
            languages_list.append(w)
            

#    if (len_languages)<=5:
#        pass
#    else:
#    print ('there are', len_languages, 'in languages list')
    lang_list=top_frequent(languages_list, k=3)
#    print ('length of languages list is:',len(lang_list))
    return(lang_list)


def extract_phone(i,sentence,phone_list,cv_name):
    
    phoneNumRegex = re.compile(r'\d\d\d-\d\d\d-\d\d\d\d')

    mo = phoneNumRegex.search(sentence)
    if mo  is not None:
        phone_list.append(mo.group())
    
    

    
def extract_email(i,sentence,email_list,cv_name):
    
    s = re.sub(r"\s+"," ", sentence, flags = re.I)

    email_list.append(re.findall('\S+@\S+', s))
    
    

#    return(email_list)

def make_couples(job_title_list,exp_ind,dates_index):
    # print('MAKING COUPLES...')
    # print ('exp index is:', exp_ind)
    # print ('dates-index is:', dates_index)
    
    # print ('exp index is:', exp_ind)
    # print ('dates index is:', dates_index)
    thresh= 4
    couples= []
    matched_dates=[]
    for i in exp_ind:
        
        # print('exp_ind', i)
        if len(dates_index)>0:
        
            closest_match= min(dates_index, key=lambda x:abs(x-i))
            # print ('closest_match' , closest_match)
            difference_inds= (i-closest_match)
            # print('difference_inds', difference_inds)
            if difference_inds>=0 and difference_inds<=thresh and closest_match not in matched_dates:
                # couples.append([i,closest_match,'end date'])
                couples.append([i,closest_match, 'start_date', 'end_date',30])
                matched_dates.append(closest_match)
                # print ('closest match is negative')
                
            else:
                thresh= -4
                if  difference_inds>=thresh and closest_match not in matched_dates:
                    # print ('difference is less than threshold value')
                # couples.append([i,closest_match,'end date'])
                    couples.append([i,closest_match,'start_date', 'end_date',30])
                    matched_dates.append(closest_match)
                
                
            # print ('there is no date present')
            # couples.append([i,'no date'])
            # matched_dates.append('no date')
        else:
            
#            print ('there is no date present')
            couples.append([i,'no date', 'start_date', 'end_date',30])  # 30 days of duration by default
            matched_dates.append('no date')
            
    # print('matched all couples')
    # print ('**************************************************************')
    return couples
        
  
def extract_job_date(i,job_title_list,sentence,experience_list,experience_index,dates_index,cv_name):
#    lsm_j=spacy.load(os.path.join (models_folder,'job_model_fre'))
    doc_job=lsm_j(sentence)
            # print('job#########',sentence,'same')
    for X in doc_job.ents:
        if X.label_=='job title':
                    
#                     print('job*******************',X.text)
            job_found=X.text
            # print ('Job found and value is:',job_found)
            experience_list.append(job_found)
            experience_index.append(i)
            job_title_list.append(job_found)
            # print ('experience found is:', job_found)
                    
    date_=search_dates(sentence,languages=['fr']) 
    if str(type(date_))!="<class 'NoneType'>":
        # print ('date found')
        # print ('date is:', date_)
        dates_index.append(i)
        
        
    
    

def extract_all_features(all_paras,cv_name):
    # print ('=============================================================')
 
    count=0
    job_found=''
    job_title_list=[]
    dates_list=[]
    experience_list=[]
    experience_index=[]
    skills_list=[]
    qualifications_list=[]
    languages_list=[]
    locations_list=[]
    all_sentences=[]
    dates_index=[]
    email_list=[]
    phone_list=[]
    all_text=all_paras.paragraphs
    
    
    
    
    # print ('length of all para is:', len(all_text))
    for i,para in enumerate(all_text):
        # print ('value of i is:',i)
        temp = []

        sentence=para.text.lower().strip()
        
            # if sentence!="":
        sentence=sentence.strip().lower()
        # sentence = re.sub(r'[.|,|:|\']',' ', sentence) #filter out all the commas, periods, and appostrophes using rege
        # sentence=re.sub('[^A-Za-z0-9À-ÿ]+', ' ', sentence)
        # sentence= sentence.decode(encoding)
        # sentence=re.sub("[\u0300-\u036f]", "", sentence)
        sentence=unidecode.unidecode(sentence)
        
        #remove multiple spaces
        sentence = re.sub(r"\s+"," ", sentence, flags = re.I)
        
        #remove spaces from start and end
        sentence = re.sub(r"^\s+", "", sentence)
        
        # remove single characters 
        sentence = re.sub(r"^\s+", "", sentence)
        all_sentences.append(sentence)
            
            # print('*****************',sentence)
        
        extract_job_date(i,job_title_list,sentence,experience_list,experience_index,dates_index,cv_name)
        # print ('job and date extracted')
        q_list=extract_qualifications(i,sentence,qualifications_list,cv_name)
        # print ('qualification extracted')
        s_list=extract_skills(i,sentence,skills_list,cv_name)
        # print ('skills extracted')
        
        lang_list=extract_languages(i,sentence,languages_list,cv_name)  #language list
        # print ('language extracted')
        loc_list=extract_locations(i,sentence,locations_list,cv_name)  #location list
        # print ('locations extracted')
        extract_phone(i,sentence,phone_list,cv_name)  #location list
        extract_email(i,sentence,email_list,cv_name)  #location list
        
        
    couples= make_couples(job_title_list,experience_index,dates_index)
    
                      
      
        
    # couples= make_couples(experience_index,dates_index)
    # print(couples)
    return couples, experience_list, experience_index,dates_index,all_sentences, q_list, s_list, loc_list, lang_list, email_list, phone_list
                        


def make_final_list(dir_cvs):    
    
    dircvs = [join(dir_cvs, f) for f in listdir(dir_cvs) if isfile(join(dir_cvs, f))]
    alltext = [] 
    cvs_list= []
    for cv in dircvs:
        print ('**************************************************************')
        print('SCANNING NEW CV')
        cvs_dict= dict()
        flag_move=''
        move=''
        name=os.path.basename(cv)
        name=os.path.splitext(name)[0]
        print ('name of cv is:', name)
        # print(cv)
        paras_list=convert_to_paras(cv)
        
         
        if paras_list!="":
            couples,experience_list, experience_index,dates_index, sent_list, qualification_list, skills_list, locations_list, languages_list,email_list, phone_list= extract_all_features(paras_list,name)
            # qualification=extract_qualification(sent_list,name)
            
            cvs_dict['name'] = name
            cvs_dict['text paras'] = sent_list
            cvs_dict['couples'] = couples
            cvs_dict['qualification']=qualification_list
            cvs_dict['skills']=skills_list
            cvs_dict['location']=locations_list            
            cvs_dict['languages']=languages_list
            cvs_dict['phone_numbers']=phone_list
            cvs_dict['emails']=email_list
            cvs_dict['experience']=experience_list
            cvs_dict['experience_index']=experience_index
            cvs_dict['date index']=dates_index
            cvs_dict['recent_job']='no job'
            cvs_dict['recent_job_experience']=0
            cvs_dict['total_experience']=0
            # cvs_dict['duration']='2 years'
            
            cvs_list.append(cvs_dict)
    return cvs_list

#==============================================================================
    #     END OF make_final_list function
#==============================================================================

data_path=os.path.join(os.path.abspath(os.path.join(__file__,"../../")),'data')
cvs_folder=os.path.join(data_path,'cvs')
models_folder=os.path.join(data_path,'models')
results_folder=os.path.join(data_path,'results')



nlp = spacy.load ('fr_core_news_lg')
lsm_j=spacy.load(os.path.join (models_folder,'job_model_fre'))
lsm_s=spacy.load(os.path.join (models_folder,'skill_model_fre'))
lsm_j=spacy.load(os.path.join (models_folder,'job_model_fre'))
lsm_q=spacy.load(os.path.join (models_folder,'Qualification_model_fre'))
location_txt=models_folder+'\locations.txt'
language_txt=models_folder+'\languages.txt'

main_dictionary={}


#========================================================================
#========================================================================


# -*- coding: utf-8 -*-
"""
Created on Fri Feb 26 14:04:20 2021

@author: KHC
"""


#========================================================================
#========================================================================


with open (location_txt, "r",encoding="utf8") as myfile:
    data=myfile.read()
                
l_list=data.split ('\n')
l_list=[l.lower() for l in l_list if l!=""]
    
    

with open (language_txt, "r",encoding="utf8") as myfile:
    data=myfile.read()
                
lang_list=data.split ('\n')
lang_list=[lang.lower() for lang in lang_list]
    

cvs_list= make_final_list(cvs_folder)
custom_stop_words=['a']
today = datetime.today()
datem = datetime(today.year, today.month, 1)

for cv_data in cvs_list:
    # print ('*********************************************************')
    print ('cv name  is:', cv_data['name'])
    # print ('*********************************************************')
    paras= cv_data['text paras']
    
    # print('paras of cv data are:',cv_data['text paras'])
    for couple in cv_data['couples']:
        

        ind1= couple[0] 
        ind2= couple[1]  
        # print ('value of ind2 is:', ind2)
        
        
        if isinstance(ind2, int):  # check if ind2 is an integer
                   
            couple[0]=paras[ind1]  # this is experience 
            couple[1]=paras[ind2]  # this is for date
            words_list = couple[1].split()

            # Using lambda expression filter the data
            sent = ' '.join((filter(lambda val: val not in custom_stop_words, words_list)))

            date_list=search_dates(sent,languages=['fr','es']) 
#            print ('date list is:', date_list)
            if str(type(date_list))!="<class 'NoneType'>":
                starting_date=date_list[0][1]
#                print ('starting date is:',starting_date)
                
                if len (date_list)>1:
                    ending_date=date_list[1][1]
#                    print('ending date is:', ending_date)
                    d=ending_date-starting_date
                    d=d.days
                   
                else:
                    if 'a ce jour' in couple[1] or 'jour' in couple[1] or 'till now' in couple[1]:
                        ending_date=datem 
                        d=ending_date-starting_date
                        d=d.days
                        # d = divmod(ending_date-starting_date,86400)  # days
#                        print ('difference in days is:',d)
                    else:
                        ending_date='end date not given'
                        d = 30 # days
#                        print ('difference in days is:',d)
                    
                    
                couple[2]=starting_date
                couple[3]=ending_date
                couple[4]=d

               
            else:
                couple[0]=paras[ind1]
                # print('couple without date  is:', couple)
                # print('paras without date  is:', paras)
                
                couple[1]='no date attached'
 
with open(os.path.join(results_folder,'ranking_list.pickle'), 'wb') as handle:
    pickle.dump(cvs_list, handle, protocol=pickle.HIGHEST_PROTOCOL)       
#df = pd.DataFrame(cvs_list) 

#df.to_csv(os.path.join(results_folder,'ranking.csv'))




