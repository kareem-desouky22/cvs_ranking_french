       # -*- coding: utf-8 -*-
"""
Created on Sun Jan 10 18:54:22 2021

@author: KHC
"""
# -*- coding: utf-8 -*-
"""
Created on Sun Jan  3 18:42:34 2021

@author: KHC
"""
# -*- coding: utf-8 -*-
"""
Created on Fri Jan  1 00:08:47 2021

@author: KHC
"""
# import fitz
import os
from os import listdir
import spacy
from os.path import isfile, join
import docx
import re
#import datetime
from datetime import datetime
from dateparser.search import search_dates
import pickle
from nltk.stem import WordNetLemmatizer
import unidecode
from docx import Document
import docx2txt
#from calendar import month_name,month_abbr
from utils_cvs import top_frequent
lemmatizer = WordNetLemmatizer()
import pandas as pd

# this function will take a single CV and will return a list of sentences extracted
#from that CV. It will extract paragraphs as first step and then will convert 
# paragraphs into sentences and then will make a list of sentences. 

encoding = "utf-8"

data_path=os.path.join(os.path.abspath(os.path.join(__file__,"../../")),'data')
cvs_folder=os.path.join(data_path,'cvs_2')
models_folder=os.path.join(data_path,'models')
results_folder=os.path.join(data_path,'results')

nlp = spacy.load ('fr_core_news_lg')
lsm_j=spacy.load(os.path.join (models_folder,'job_model_fre'))
lsm_s=spacy.load(os.path.join (models_folder,'skill_model_fre'))
lsm_q=spacy.load(os.path.join (models_folder,'Qualification_model_fre'))
#location_txt=models_folder+'\locations.txt'
#language_txt=models_folder+'\languages.txt'

main_dictionary={}
#========================================================================
#========================================================================

#with open (location_txt, "r",encoding="utf8") as myfile:
#    data=myfile.read()
#
#
##todo create a function to perform this action                 
#l_list=data.split ('\n')
#l_list=[l.lower() for l in l_list if l!=""]     
#
#with open (language_txt, "r",encoding="utf8") as myfile:
#    data=myfile.read()                
#lang_list=data.split ('\n')
#lang_list=[lang.lower() for lang in lang_list] 

def convert_to_paras(filename):

    try:
        doc=docx.Document(filename)
        all_paras = doc
        if len(doc.paragraphs)<=1:
            my_text = docx2txt.process(filename)
            d=Document()
            d.add_paragraph(my_text)    
            all_paras = Document() 
            for s in d.paragraphs:
                str_list = s.text.split("\n")
                str_list[:] = [x.strip() for x in str_list if x]
                new_para = ""
                prev = str_list[0]
                for i in range(1, len(str_list)):
                    if new_para =="":
                        new_para = new_para + prev
                    else:
                        new_para = new_para + "\n" + prev
                    prev = str_list[i]
                    all_paras.add_paragraph(prev)
        return all_paras
    except:
         return "" 
        
       
def extract_qualifications(i,sentence,qualifications_list,cv_name):
    doc_qual=lsm_q(sentence)
    for X in doc_qual.ents:
        if X.label_=="qualification":
            contains_digit = any(map(str.isdigit, X.text))
            if contains_digit or ':' in X.text:
                pass
            else:            
                qualifications_list.append(X.text)
              
    qualifications_list=top_frequent(qualifications_list, k=15)
    return(qualifications_list)
            
def extract_skills(i,sentence,skills_list,cv_name):  # why using i ?
    doc_skills=lsm_s(sentence)
    for X in doc_skills.ents:
        if X.label_=="skills":
            

            contains_digit = any(map(str.isdigit, X.text))
            if contains_digit:            
                
                print ('skill is skipped and value is:', X.text)
                
            elif ':' in X.text:
                skill=X.text.split(":")[1]
                print ('skill is:', skill)
                
            else:
                skills_list.append(X.text)
    skills_list=top_frequent(skills_list, k=15)
    return(skills_list)
        
            
def extract_locations(i,sentence,locations_list,cv_name):
#    global nlp,l_list
    doc_all=nlp(sentence)
    for X in doc_all.ents:
        if (X.label_=='GPE'):
            locations_list.append(X.text)            
#    words_list=sentence.split(' ')
#    for w in words_list:
#        if w in l_list:
#            locations_list.append(w)
    loc_list=top_frequent(locations_list, k=3)
    return(loc_list)
            
def extract_languages(i,sentence,languages_list,cv_name):
#    global nlp,lang_list
    doc_all=nlp(sentence)
    for X in doc_all.ents:
        if (X.label_=='LANGUAGE'):
            languages_list.append(X.text)
#    words_list=sentence.split(' ')
#    print (sss)
#    for w in words_list:
#        if w in lang_list:
#            languages_list.append(w)
    lang_list=top_frequent(languages_list, k=3)
    return(lang_list)


def extract_phone(i,sentence,phone_list,cv_name):    
    phoneNumRegex = re.compile(r'\d\d\d-\d\d\d-\d\d\d\d')
    mo = phoneNumRegex.search(sentence)
    if mo  is not None:
        phone_list.append(mo.group())

def extract_email(i,sentence,email_list,cv_name):    
    s = re.sub(r"\s+"," ", sentence, flags = re.I)
    email_list.append(re.findall('\S+@\S+', s))


def make_couples(job_title_list,job_ind,jobs_dates_index):
    print ('job_title_list is:', job_title_list)
    print ('job inds are:', job_ind)
    print ('date inds are:', jobs_dates_index)
    thresh= 4
    couples= []
    matched_dates=[]
    for i in job_ind:
        if len(jobs_dates_index)>0:
            closest_match= min(jobs_dates_index, key=lambda x:abs(x-i))
            difference_inds= (i-closest_match)
            if difference_inds>=0 and difference_inds<=thresh and closest_match not in matched_dates:
                print('')
                print ('closest match is:', closest_match)
                print ('couple list value is:', couples)
                couples.append([i,closest_match, 'start_date', 'end_date',30])
                matched_dates.append(closest_match)
              
            else:
                thresh= -4
                if  difference_inds>=thresh and closest_match not in matched_dates:
                    couples.append([i,closest_match,'start_date', 'end_date',30])
                    matched_dates.append(closest_match)
                    print ('closest match is:', closest_match)
                    print ('couple list value is:', couples)
        else:
            couples.append([i,'no date', 'start_date', 'end_date',30])  # 30 days of duration by default
            matched_dates.append('no date')

    return couples
        
  
def extract_job_date(i,sentence,jobs_titles_list,jobs_titles_index_list,jobs_dates_index,cv_name):
    doc_job=lsm_j(sentence)
    for X in doc_job.ents:
        if X.label_=='job title':
#            job_found=X.text
            job_found=X.ents
#            print ('job found is:',job_found)
            jobs_titles_list.append(job_found)
            jobs_titles_index_list.append(i)
            

    date_=search_dates(sentence,languages=['fr']) 
    if str(type(date_))!="<class 'NoneType'>":
        jobs_dates_index.append(i)
        

def extract_all_features(all_paras,cv_name):

    jobs_titles_list=[]
    jobs_titles_index_list=[]
    skills_list=[]
    qualifications_list=[]
    languages_list=[]
    locations_list=[]
    all_sentences=[]
    jobs_dates_index=[]
    email_list=[]
    phone_list=[]
    all_text=all_paras.paragraphs

    for i,para in enumerate(all_text):
        sentence=para.text.lower().strip()
        sentence=sentence.strip()
        sentence=unidecode.unidecode(sentence)
        
        #remove multiple spaces
        sentence = re.sub(r"\s+"," ", sentence, flags = re.I)
        sentence = re.sub(r"^\s+", "", sentence)
        sentence = re.sub(r"^\s+", "", sentence)
        all_sentences.append(sentence)   # Whats the use of all_sentence ?
        extract_job_date(i,sentence,jobs_titles_list,jobs_titles_index_list,jobs_dates_index,cv_name)
        q_list=extract_qualifications(i,sentence,qualifications_list,cv_name)
        s_list=extract_skills(i,sentence,skills_list,cv_name)
        extract_languages(i,sentence,languages_list,cv_name)  #language list
        extract_locations(i,sentence,locations_list,cv_name)  #location list
        extract_phone(i,sentence,phone_list,cv_name)  #location list
        extract_email(i,sentence,email_list,cv_name)  #location list
    couples= make_couples(jobs_titles_list,jobs_titles_index_list,jobs_dates_index)
    return couples, jobs_titles_list, jobs_titles_index_list,jobs_dates_index,all_sentences, q_list, s_list, locations_list, languages_list, email_list, phone_list
                        


def make_final_list(dir_cvs):    
    
    dir_cvs = [join(dir_cvs, f) for f in listdir(dir_cvs) if isfile(join(dir_cvs, f))]
#    alltext = [] 
    cvs_list= []
    for cv in dir_cvs:
        print ('**************************************************************')
        print('SCANNING NEW CV')
        cvs_dict= dict()
        name=os.path.basename(cv)
        name=os.path.splitext(name)[0]
        print ('name of cv is:', name)
        # print(cv)
        paras_list=convert_to_paras(cv)
        
         
        if paras_list!="":
            couples,jobs_titles_list, jobs_titles_index_list,jobs_dates_index, sent_list, qualification_list, skills_list, locations_list, languages_list,email_list, phone_list= extract_all_features(paras_list,name)
            # qualification=extract_qualification(sent_list,name)
            
            cvs_dict['name_cv'] = name
            cvs_dict['text paras'] = sent_list
            cvs_dict['couples'] = couples
            cvs_dict['qualification_cv']=qualification_list
            cvs_dict['skills_cv']=skills_list
            cvs_dict['location_cv']=locations_list            
            cvs_dict['languages_cv']=languages_list
            cvs_dict['phone_numbers_cv']=phone_list
            cvs_dict['emails_cv']=email_list
            cvs_dict['experience_cv']=jobs_titles_list
            cvs_dict['jobs_titles_index_list']=jobs_titles_index_list
            cvs_dict['date index']=jobs_dates_index
            cvs_dict['recent_job_cv']='no job'
            cvs_dict['recent_job_experience_cv']=0
            cvs_dict['total_experience_cv']=0
            # cvs_dict['duration']='2 years'
            
            cvs_list.append(cvs_dict)
    return cvs_list

#==============================================================================
    #     END OF make_final_list function
#==============================================================================


   


cvs_list= make_final_list(cvs_folder)
custom_stop_words=['a']
today = datetime.today()
datem = datetime(today.year, today.month, 1)

for cv_data in cvs_list:
    paras= cv_data['text paras']
    for couple in cv_data['couples']:
        ind1= couple[0] 
        ind2= couple[1]  
        if isinstance(ind2, int):  # check if ind2 is an integer                  
            couple[0]=paras[ind1]  # this is experience 
            couple[1]=paras[ind2]  # this is for date
            words_list = couple[1].split()
            sent = ' '.join((filter(lambda val: val not in custom_stop_words, words_list)))
            date_list=search_dates(sent,languages=['fr','es']) 
            if str(type(date_list))!="<class 'NoneType'>":
                starting_date=date_list[0][1]               
                if len (date_list)>1:
                    ending_date=date_list[1][1]
                    d=ending_date-starting_date
                    d=d.days                   
                else:
                    if 'a ce jour' in couple[1] or 'jour' in couple[1] or 'till now' in couple[1]:
                        ending_date=datem 
                        d=ending_date-starting_date
                        d=d.days
                    else:
                        ending_date='end date not given'
                        d = 30 # days
                couple[2]=starting_date
                couple[3]=ending_date
                couple[4]=d            
            else:
                couple[0]=paras[ind1]               
                couple[1]='no date attached'
 
#with open(os.path.join(results_folder,'ranking_list1.pickle'), 'wb') as handle:
#    pickle.dump(cvs_list, handle, protocol=pickle.HIGHEST_PROTOCOL)       
df = pd.DataFrame(cvs_list) 

df.to_csv(os.path.join(results_folder,'ranking1.csv'))




